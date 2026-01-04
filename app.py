import streamlit as st
import base64
import PyPDF2
import pandas as pd  # 👈 新增：数据分析神器
import io
from PIL import Image
from openai import OpenAI
from docx import Document

# 1. 配置 DeepSeek 客户端
client = OpenAI(
    api_key=st.secrets["DEEPSEEK_API_KEY"],
    base_url="https://api.deepseek.com"
)


# --- 工具函数 1：图片转码 ---
def encode_image(img_file):
    if img_file is None: return None
    bytes_data = img_file.getvalue()
    return base64.b64encode(bytes_data).decode('utf-8')


# --- 工具函数 2：PDF 转文字 ---
def extract_text_from_pdf(pdf_file):
    if pdf_file is None: return ""
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text


# --- 工具函数 3：Excel 转文字 (新增！核心功能) ---
def get_excel_data(excel_file):
    """
    读取 Excel，在界面展示，并返回给 AI 可读的文本
    """
    try:
        # 使用 pandas 读取
        df = pd.read_excel(excel_file)

        # 1. 在网页上显示预览 (只显示前 50 行，防止太卡)
        with st.expander("📊 点击查看表格数据预览", expanded=True):
            st.dataframe(df)
            st.caption(f"共检测到 {df.shape[0]} 行，{df.shape[1]} 列数据。")

        # 2. 转换成 AI 能看懂的 CSV 格式字符串
        # 为了节省 Token，如果表格太大，我们只截取前 100 行给 AI
        if len(df) > 100:
            csv_text = df.head(100).to_csv(index=False)
            warning = "\n[系统注：数据量较大，已截取前100行供分析]"
            return csv_text + warning
        else:
            return df.to_csv(index=False)

    except Exception as e:
        st.error(f"Excel 读取失败: {e}")
        return ""


# --- 工具函数 4：生成 Word 文档 ---
def create_docx(content):
    doc = Document()
    doc.add_heading('AI 项目分析报告', 0)
    doc.add_paragraph(content)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# --- 侧边栏设置 ---
def setup_sidebar():
    st.sidebar.title("🤖 科技项目全能助手")

    # 任务模式
    task_type = st.sidebar.selectbox(
        "选择任务模式 🛠️",
        ["通用助手", "Excel数据分析", "公文润色/仿写", "项目申报书撰写"],
        index=0
    )

    # 自动匹配人设
    prompts = {
        "通用助手": "你是一名资深的科技项目专家。",
        "Excel数据分析": "你是一名精通数据分析的商业智能(BI)专家。请根据用户上传的表格数据（CSV格式），进行逻辑分析、计算关键指标，并发现数据背后的趋势或问题。回答要基于数据，严谨客观。",
        "公文润色/仿写": "你是一名公文写作专家，语言庄重、严谨。",
        "项目申报书撰写": "你是一名科技项目申报顾问，擅长撰写逻辑清晰的申报材料。"
    }

    default_prompt = prompts[task_type]
    system_prompt = st.sidebar.text_area("系统人设", value=default_prompt, height=150)

    if "messages" in st.session_state:
        st.session_state.messages[0] = {"role": "system", "content": system_prompt}

    # 模型配置
    model_name = st.sidebar.selectbox("选择模型", ["deepseek-chat", "deepseek-coder"], index=0)
    temperature = st.sidebar.slider("创造力", 0.0, 1.0, 0.2)

    # 文件上传 (新增 xlsx 支持)
    uploaded_file = st.sidebar.file_uploader(
        "上传资料 (PDF/Excel/图片)",
        type=["jpg", "png", "pdf", "xlsx"]  # 👈 加入了 xlsx
    )

    if uploaded_file:
        file_type = uploaded_file.name.split('.')[-1].lower()
        st.sidebar.success(f"📂 已加载: {uploaded_file.name} ({file_type})")

    if st.sidebar.button("🗑️ 清空对话"):
        st.session_state.messages = [{"role": "system", "content": system_prompt}]
        st.rerun()

    return model_name, temperature, uploaded_file


# --- 主程序 ---
def main():
    st.title("📊 四川平高-智能数据分析台")

    if "messages" not in st.session_state:
        st.session_state.messages = [{"role": "system", "content": "你是一名专家。"}]

    model_name, temperature, uploaded_file = setup_sidebar()

    for msg in st.session_state.messages:
        if msg["role"] == "system": continue
        with st.chat_message(msg["role"]):
            st.write(msg["content"])

    if user_input := st.chat_input("请输入指令 (例如：帮我分析下哪个项目预算最高？)..."):

        with st.chat_message("user"):
            st.write(user_input)
        st.session_state.messages.append({"role": "user", "content": user_input})

        with st.chat_message("assistant"):
            message_placeholder = st.empty()
            full_response = ""
            messages_to_send = st.session_state.messages.copy()

            # === 处理各种文件类型 ===
            if uploaded_file:
                file_type = uploaded_file.name.split('.')[-1].lower()

                # 1. Excel 处理逻辑
                if file_type == 'xlsx':
                    with st.spinner("正在解析 Excel 数据..."):
                        data_context = get_excel_data(uploaded_file)
                    new_content = f"【数据表格内容(CSV格式)】：\n{data_context}\n\n【用户问题】：{user_input}"
                    messages_to_send[-1]["content"] = new_content

                # 2. PDF 处理逻辑
                elif file_type == 'pdf':
                    with st.spinner("正在阅读 PDF..."):
                        pdf_text = extract_text_from_pdf(uploaded_file)
                    new_content = f"【背景资料】：\n{pdf_text[:20000]}\n\n【用户指令】：{user_input}"
                    messages_to_send[-1]["content"] = new_content

                # 3. 图片处理逻辑
                elif file_type in ['jpg', 'jpeg', 'png']:
                    new_content = user_input + f"\n[系统注：用户上传了图片 '{uploaded_file.name}']"
                    messages_to_send[-1]["content"] = new_content

            # === 发送请求 ===
            try:
                response = client.chat.completions.create(
                    model=model_name,
                    messages=messages_to_send,
                    temperature=temperature,
                    stream=True
                )

                for chunk in response:
                    content = chunk.choices[0].delta.content or ""
                    full_response += content
                    message_placeholder.markdown(full_response + "▌")
                message_placeholder.markdown(full_response)

                st.session_state.messages.append({"role": "assistant", "content": full_response})

                # Word 下载按钮
                if full_response:
                    docx_file = create_docx(full_response)
                    st.download_button(
                        label="📥 将分析结果下载为 Word",
                        data=docx_file,
                        file_name="Analysis_Report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            except Exception as e:
                st.error(f"❌ 发生错误: {e}")


if __name__ == "__main__":
    main()